"""统计报告渲染 backend.stats.report 单元测试（HTML 与 Word）。"""
from docx import Document

from backend.stats.report import build_stats_docx, build_stats_html


def _sample_stats(**overrides) -> dict:
    """构造一份 analyze_batch 结构的样例统计数据。"""
    data = {
        "summary": {
            "total_records": 3,
            "question_count": 2,
            "avg_score": 71.3,
            "max_score": 100.0,
            "min_score": 40.0,
        },
        "questions": [
            {
                "question_no": 1, "count": 2, "avg_score": 80.0,
                "max_score": 100.0, "min_score": 60.0,
                "pass_count": 2, "pass_rate": 100.0, "unanswered_count": 0,
            },
            {
                "question_no": 2, "count": 3, "avg_score": None,
                "max_score": None, "min_score": None,
                "pass_count": 1, "pass_rate": 66.7, "unanswered_count": 2,
            },
        ],
        "categories": [
            {"error_category": "grammar", "count": 2, "avg_score": 80.0},
            {"error_category": "math", "count": 1, "avg_score": 54.0},
        ],
    }
    data.update(overrides)
    return data


def test_build_stats_html_contains_summary_and_table_headers():
    """HTML 报告：含总体概览数值、每题表头与错因表头。"""
    html = build_stats_html(_sample_stats())
    assert "批改统计报告" in html
    assert "总批改记录：3" in html
    assert "题目数：2" in html
    assert "平均分：71.3" in html
    assert "最高分：100.0" in html
    assert "最低分：40.0" in html
    for header in ("题号", "人数", "平均", "最高", "最低", "及格率", "未作答"):
        assert f"<th>{header}</th>" in html
    for header in ("错因", "人数", "平均分"):
        assert f"<th>{header}</th>" in html


def test_build_stats_html_none_shows_dash_and_pass_rate_integer_percent():
    """None 显示 '--'，pass_rate 渲染为整数百分比。"""
    html = build_stats_html(_sample_stats())
    assert "--" in html
    assert "100%" in html
    assert "67%" in html


def test_build_stats_html_empty_data_renders_placeholder():
    """空统计结构：概览显示 '--'，表格给占位行。"""
    html = build_stats_html({"summary": {}, "questions": [], "categories": []})
    assert "总批改记录：--" in html
    assert "题目数：--" in html
    assert "暂无数据" in html


def test_build_stats_docx_writes_openable_file_and_returns_path(tmp_path):
    """Word 报告：写入 tmp_path，文件存在且可读，返回文件路径。"""
    out_path = tmp_path / "stats.docx"
    returned = build_stats_docx(_sample_stats(), str(out_path))
    assert returned == str(out_path)
    assert out_path.exists()
    assert out_path.stat().st_size > 0

    doc = Document(str(out_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "批改统计报告" in text
    assert "总批改记录：3" in text
    assert "题目数：2" in text

    assert len(doc.tables) == 2
    question_headers = [cell.text for cell in doc.tables[0].rows[0].cells]
    assert question_headers == ["题号", "人数", "平均", "最高", "最低", "及格率", "未作答"]
    category_headers = [cell.text for cell in doc.tables[1].rows[0].cells]
    assert category_headers == ["错因", "人数", "平均分"]
