"""报告生成模块：将批改结果渲染为独立 HTML 或 Word 报告。"""
from html import escape

# 评分方式显示文案，键为 grade_answer 返回的 method
_METHOD_LABELS = {
    "online": "在线 DeepSeek 精排",
    "offline": "离线向量相似度",
}

# HTML 报告内嵌样式，使用浅色打印友好的配色
_STYLE = """
body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "Microsoft YaHei", sans-serif; margin: 40px auto; max-width: 800px; padding: 0 20px; color: #1c1c1e; line-height: 1.6; }
h1 { color: #007aff; border-bottom: 2px solid #007aff; padding-bottom: 10px; }
h2 { color: #005bb5; margin-top: 30px; }
.meta { background: #f5f5f7; padding: 16px 20px; border-radius: 10px; margin: 20px 0; }
.meta p { margin: 6px 0; }
.score { font-size: 28px; font-weight: bold; color: #34c759; }
.content { background: #ffffff; border: 1px solid #e5e5ea; border-radius: 10px; padding: 16px 20px; white-space: pre-wrap; }
@media print { body { margin: 0; } .meta { background: #fafafa; } }
"""


def _yes_no(value) -> str:
    """将布尔/类布尔值转换为中文是/否文案。"""
    return "是" if value else "否"


def _method_label(method) -> str:
    """将 method 值映射为展示文案，未知值原样返回。"""
    return _METHOD_LABELS.get(method, method or "未知")


def build_report_html(data: dict) -> str:
    """将批改结果渲染为独立 HTML 报告字符串，可直接在浏览器打开或打印。

    data 键约定：work_content / answer_content / score / method / degraded /
    routed / generated_at（ISO 字符串）。
    """
    score = data.get("score", "--")
    generated_at = data.get("generated_at", "")
    work_content = escape(str(data.get("work_content") or ""))
    answer_content = escape(str(data.get("answer_content") or ""))
    body = f"""
<h1>作业批改报告</h1>
<p>批改时间：{escape(str(generated_at))}</p>
<div class="meta">
<p>作业评分（百分制）：<span class="score">{score}</span></p>
<p>评分方式：{_method_label(data.get("method"))}</p>
<p>是否降级：{_yes_no(data.get("degraded"))}</p>
<p>是否路由精排：{_yes_no(data.get("routed"))}</p>
</div>
<h2>学生作业内容</h2>
<div class="content">{work_content}</div>
<h2>参考答案内容</h2>
<div class="content">{answer_content}</div>
"""
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>作业批改报告</title>
<style>{_STYLE}</style>
</head>
<body>
{body}
</body>
</html>"""


def build_report_docx(data: dict, out_path: str) -> str:
    """用 python-docx 生成 Word 报告写入 out_path，返回文件路径。

    python-docx 采用函数内懒导入，便于离线测试时 mock 依赖。
    """
    from docx import Document

    document = Document()
    document.add_heading("作业批改报告", level=0)
    document.add_paragraph(f"批改时间：{data.get('generated_at', '')}")
    document.add_paragraph(f"作业评分（百分制）：{data.get('score', '--')}")
    document.add_paragraph(f"评分方式：{_method_label(data.get('method'))}")
    document.add_paragraph(f"是否降级：{_yes_no(data.get('degraded'))}")
    document.add_paragraph(f"是否路由精排：{_yes_no(data.get('routed'))}")
    document.add_heading("学生作业内容", level=1)
    document.add_paragraph(str(data.get("work_content") or ""))
    document.add_heading("参考答案内容", level=1)
    document.add_paragraph(str(data.get("answer_content") or ""))
    document.save(out_path)
    return out_path
