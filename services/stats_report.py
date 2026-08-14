"""统计报告渲染：将 analyze_batch 结果渲染为独立 HTML 或 Word 报告。"""
from __future__ import annotations

from html import escape

# HTML 报告内嵌样式，浅色打印友好配色，与 report.py 风格一致
_STYLE = """
body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "Microsoft YaHei", sans-serif; margin: 40px auto; max-width: 800px; padding: 0 20px; color: #1c1c1e; line-height: 1.6; }
h1 { color: #007aff; border-bottom: 2px solid #007aff; padding-bottom: 10px; }
h2 { color: #005bb5; margin-top: 30px; }
.meta { background: #f5f5f7; padding: 16px 20px; border-radius: 10px; margin: 20px 0; }
.meta p { margin: 6px 0; }
table { width: 100%; border-collapse: collapse; margin: 12px 0 24px; }
th, td { border: 1px solid #e5e5ea; padding: 8px 10px; text-align: left; font-size: 14px; }
th { background: #f5f5f7; font-weight: 600; }
.empty { text-align: center; color: #8e8e93; }
@media print { body { margin: 0; } .meta { background: #fafafa; } }
"""

# 每题统计表列标题
_QUESTION_HEADERS = ("题号", "人数", "平均", "最高", "最低", "及格率", "未作答")
# 错因分布表列标题
_CATEGORY_HEADERS = ("错因", "人数", "平均分")


def _fmt(value) -> str:
    """数值格式化：None/缺失显示 '--'，其余原样显示。"""
    return "--" if value is None else str(value)


def _fmt_pct(value) -> str:
    """百分比显示为整数；None/缺失显示 '--'。"""
    if value is None:
        return "--"
    return f"{round(float(value))}%"


def _question_rows(questions: list) -> str:
    """渲染每题统计表行 HTML，空数据给占位行。"""
    if not questions:
        return f'<tr><td colspan="{len(_QUESTION_HEADERS)}" class="empty">暂无数据</td></tr>'
    return "\n".join(
        f"<tr><td>{_fmt(q.get('question_no'))}</td>"
        f"<td>{_fmt(q.get('count'))}</td>"
        f"<td>{_fmt(q.get('avg_score'))}</td>"
        f"<td>{_fmt(q.get('max_score'))}</td>"
        f"<td>{_fmt(q.get('min_score'))}</td>"
        f"<td>{_fmt_pct(q.get('pass_rate'))}</td>"
        f"<td>{_fmt(q.get('unanswered_count'))}</td></tr>"
        for q in questions
    )


def _category_rows(categories: list) -> str:
    """渲染错因分布表行 HTML，空数据给占位行。"""
    if not categories:
        return f'<tr><td colspan="{len(_CATEGORY_HEADERS)}" class="empty">暂无数据</td></tr>'
    return "\n".join(
        f"<tr><td>{escape(str(c.get('error_category')) or '--')}</td>"
        f"<td>{_fmt(c.get('count'))}</td>"
        f"<td>{_fmt(c.get('avg_score'))}</td></tr>"
        for c in categories
    )


def build_stats_html(data: dict) -> str:
    """将 analyze_batch 统计结构渲染为独立 HTML 报告字符串。"""
    summary = data.get("summary") or {}
    questions = data.get("questions") or []
    categories = data.get("categories") or []

    headers_q = "".join(f"<th>{h}</th>" for h in _QUESTION_HEADERS)
    headers_c = "".join(f"<th>{h}</th>" for h in _CATEGORY_HEADERS)
    body = f"""
<h1>批改统计报告</h1>
<div class="meta">
<p>总批改记录：{_fmt(summary.get('total_records'))}</p>
<p>题目数：{_fmt(summary.get('question_count'))}</p>
<p>平均分：{_fmt(summary.get('avg_score'))}</p>
<p>最高分：{_fmt(summary.get('max_score'))}</p>
<p>最低分：{_fmt(summary.get('min_score'))}</p>
</div>
<h2>每题统计</h2>
<table>
<thead><tr>{headers_q}</tr></thead>
<tbody>
{_question_rows(questions)}
</tbody>
</table>
<h2>错因分布</h2>
<table>
<thead><tr>{headers_c}</tr></thead>
<tbody>
{_category_rows(categories)}
</tbody>
</table>
"""
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>批改统计报告</title>
<style>{_STYLE}</style>
</head>
<body>
{body}
</body>
</html>"""


def build_stats_docx(data: dict, out_path: str) -> str:
    """用 python-docx 生成 Word 统计报告写入 out_path，返回文件路径。

    python-docx 采用函数内懒导入，便于离线测试时 mock 依赖。
    """
    from docx import Document

    summary = data.get("summary") or {}
    questions = data.get("questions") or []
    categories = data.get("categories") or []

    document = Document()
    document.add_heading("批改统计报告", level=0)
    document.add_paragraph(f"总批改记录：{_fmt(summary.get('total_records'))}")
    document.add_paragraph(f"题目数：{_fmt(summary.get('question_count'))}")
    document.add_paragraph(f"平均分：{_fmt(summary.get('avg_score'))}")
    document.add_paragraph(f"最高分：{_fmt(summary.get('max_score'))}")
    document.add_paragraph(f"最低分：{_fmt(summary.get('min_score'))}")

    document.add_heading("每题统计", level=1)
    _fill_docx_table(document, questions, _QUESTION_HEADERS, _row_values_of_question)

    document.add_heading("错因分布", level=1)
    _fill_docx_table(document, categories, _CATEGORY_HEADERS, _row_values_of_category)

    document.save(out_path)
    return out_path


def _fill_docx_table(document, rows: list, headers: tuple, row_values) -> None:
    """向 docx 文档追加带表头的统计表并逐行填充。"""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for item in rows:
        for cell, value in zip(table.add_row().cells, row_values(item)):
            cell.text = value


def _row_values_of_question(q: dict) -> tuple:
    """把每题统计项映射为表行值。"""
    return (
        _fmt(q.get("question_no")),
        _fmt(q.get("count")),
        _fmt(q.get("avg_score")),
        _fmt(q.get("max_score")),
        _fmt(q.get("min_score")),
        _fmt_pct(q.get("pass_rate")),
        _fmt(q.get("unanswered_count")),
    )


def _row_values_of_category(c: dict) -> tuple:
    """把错因分布项映射为表行值。"""
    return (
        _fmt(c.get("error_category")),
        _fmt(c.get("count")),
        _fmt(c.get("avg_score")),
    )
